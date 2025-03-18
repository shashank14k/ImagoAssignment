import pandas as pd
import numpy as np
from scipy import stats
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
import argparse
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
import os
from matplotlib.gridspec import GridSpec


def parse_arguments():
    parser = argparse.ArgumentParser(description='Dataset Analysis Script')
    parser.add_argument('--data', type=str, required=True, help='Path to dataset CSV file')
    parser.add_argument('--report', type=str, default='analysis_report.docx', help='Output report path')
    parser.add_argument('--alpha', type=float, default=0.05, help='Significance level for normality tests')
    parser.add_argument('--top_n', type=int, default=10, help='Number of top items to display in analyses')
    parser.add_argument('--pca_var', type=float, default=0.95, help='Variance to retain in PCA')
    return parser.parse_args()


def analyze_target_column(df, target_col, alpha=0.05):
    """
    Analyze the target column for outliers and normality.

    Parameters:
    df (pandas.DataFrame): DataFrame containing the data
    target_col (str): Name of the target column
    alpha (float): Significance level for normality test (default: 0.05)

    Returns:
    dict: Dictionary containing analysis results
    """
    target_data = df[target_col].dropna()

    # Basic statistics
    stats_dict = {
        'count': len(target_data),
        'mean': target_data.mean(),
        'median': target_data.median(),
        'std': target_data.std(),
        'min': target_data.min(),
        'max': target_data.max()
    }

    # Check for outliers using IQR method
    q1 = target_data.quantile(0.25)
    q3 = target_data.quantile(0.75)
    iqr = q3 - q1
    lower_bound = q1 - 1.5 * iqr
    upper_bound = q3 + 1.5 * iqr

    outliers = target_data[(target_data < lower_bound) | (target_data > upper_bound)]
    outlier_count = len(outliers)
    outlier_percentage = (outlier_count / len(target_data)) * 100

    # Check normality using Shapiro-Wilk test
    _, p_value = stats.shapiro(target_data)
    is_normal = p_value > alpha

    # Calculate skewness and kurtosis
    skewness = target_data.skew()
    kurtosis = target_data.kurtosis()

    return {
        'statistics': stats_dict,
        'outliers': {
            'count': outlier_count,
            'percentage': outlier_percentage,
            'lower_bound': lower_bound,
            'upper_bound': upper_bound,
            'values': outliers.tolist() if len(outliers) < 10 else outliers.head(10).tolist()
        },
        'normality': {
            'p_value': p_value,
            'is_normal': is_normal,
            'interpretation': 'Normal' if is_normal else 'Non-normal',
            'skewness': skewness,
            'kurtosis': kurtosis
        }
    }


def main():
    args = parse_arguments()

    df = pd.read_csv(args.data)
    id_col = df.columns[0]
    target_col = df.columns[-1]
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    feature_cols = [col for col in numeric_cols if col not in [id_col, target_col]]

    doc = Document()
    doc.add_heading('Data Analysis Report', 0)

    doc.add_heading('1. Null Values Analysis', level=1)
    null_counts = df[feature_cols].isnull().sum().sort_values(ascending=False)
    total_null = null_counts.sum()
    if total_null > 0:
        doc.add_paragraph(f'Total null values: {total_null}')
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Column'
        hdr_cells[1].text = 'Null Count'
        for col, count in null_counts.head(args.top_n).items():
            row_cells = table.add_row().cells
            row_cells[0].text = col
            row_cells[1].text = str(count)
    else:
        doc.add_paragraph('No null values found.')

    doc.add_heading('2. Normality Assessment', level=1)
    non_normal_cols = []
    normal_cols = []
    normality_dict = {}  # Store normality results for later use

    for col in feature_cols:
        data = df[col].dropna()
        if len(data) > 3:
            _, p_value = stats.shapiro(data)
            if p_value < args.alpha:
                non_normal_cols.append(col)
                normality_dict[col] = "Non-normal"
            else:
                normal_cols.append(col)
                normality_dict[col] = "Normal"
        else:
            normality_dict[col] = "Insufficient data"

    doc.add_paragraph(f'Non-normal features: {len(non_normal_cols)}')
    if non_normal_cols:
        doc.add_paragraph(f'Top {args.top_n}: {", ".join(non_normal_cols[:args.top_n])}')
    doc.add_paragraph(f'Normal features: {len(normal_cols)}')
    if normal_cols:
        doc.add_paragraph(f'Top {args.top_n}: {", ".join(normal_cols[:args.top_n])}')

    doc.add_heading('3. Correlation Analysis', level=1)
    target_corrs = df[feature_cols].corrwith(df[target_col]).sort_values(ascending=False)
    corr_matrix = df[feature_cols].corr()
    upper_triangle = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
    feature_corrs = upper_triangle.stack().sort_values(ascending=False)

    doc.add_paragraph(f'Top {args.top_n} Feature-Target Correlations:')
    table = doc.add_table(rows=1, cols=3)  # Added a column for normality
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Correlation'
    hdr_cells[2].text = 'Distribution'  # New column for normality

    for col, val in target_corrs.head(args.top_n).items():
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = f'{val:.2f}'
        row_cells[2].text = normality_dict.get(col, "Unknown")  # Add normality information

    # Generate scatter plots for top n features vs target
    top_features = target_corrs.head(args.top_n).index.tolist()

    # Create a subplot layout for feature-target relationships
    plot_rows = (len(top_features) + 1) // 2  # Calculate rows needed (2 plots per row)
    fig = plt.figure(figsize=(12, 3 * plot_rows))
    gs = GridSpec(plot_rows, 2, figure=fig)

    for i, feature in enumerate(top_features):
        ax = fig.add_subplot(gs[i // 2, i % 2])
        sns.scatterplot(x=df[feature], y=df[target_col], alpha=0.6, ax=ax)

        # Calculate and display correlation coefficient
        corr_val = target_corrs[feature]
        ax.set_title(f'{feature} vs {target_col} (r = {corr_val:.2f})')
        ax.set_xlabel(feature)
        ax.set_ylabel(target_col)

    plt.tight_layout()
    plt.savefig('top_features_vs_target.png', dpi=300, bbox_inches='tight')
    plt.close()

    doc.add_paragraph('\nTop Features vs Target:')
    doc.add_picture('top_features_vs_target.png', width=Inches(6))

    doc.add_paragraph(f'\nTop {args.top_n} Feature-Feature Correlations:')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature 1'
    hdr_cells[1].text = 'Feature 2'
    hdr_cells[2].text = 'Correlation'
    for (f1, f2), val in feature_corrs.head(args.top_n).items():
        row_cells = table.add_row().cells
        row_cells[0].text = f1
        row_cells[1].text = f2
        row_cells[2].text = f'{val:.2f}'

    # Add correlation heatmap
    plt.figure(figsize=(12, 10))
    mask = np.triu(np.ones_like(corr_matrix, dtype=bool))
    sns.heatmap(corr_matrix, mask=mask, annot=False, cmap='coolwarm',
                vmin=-1, vmax=1, center=0, square=True, linewidths=.5)
    plt.title('Feature Correlation Heatmap')
    plt.tight_layout()
    plt.savefig('correlation_heatmap.png', dpi=300, bbox_inches='tight')
    plt.close()

    doc.add_paragraph('\nFeature Correlation Heatmap:')
    doc.add_picture('correlation_heatmap.png', width=Inches(6))

    doc.add_heading('4. Skewness Analysis', level=1)
    skewness = df[feature_cols].skew().sort_values(key=abs, ascending=False)
    skewed_cols = skewness[abs(skewness) > 1]
    doc.add_paragraph(f'Skewed features: {len(skewed_cols)}')
    if not skewed_cols.empty:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Feature'
        hdr_cells[1].text = 'Skewness'
        for col, val in skewed_cols.head(args.top_n).items():
            row_cells = table.add_row().cells
            row_cells[0].text = col
            row_cells[1].text = f'{val:.2f}'

    doc.add_heading('5. Outlier Analysis', level=1)
    outlier_counts = {}
    outlier_percentages = {}

    for col in feature_cols:
        data = df[col].dropna()
        total_values = len(data)
        if total_values > 0:
            q1 = data.quantile(0.25)
            q3 = data.quantile(0.75)
            iqr = q3 - q1
            outliers = data[(data < (q1 - 1.5 * iqr)) | (data > (q3 + 1.5 * iqr))]
            outlier_counts[col] = len(outliers)
            outlier_percentages[col] = (len(outliers) / total_values) * 100
        else:
            outlier_counts[col] = 0
            outlier_percentages[col] = 0

    outlier_percentage_series = pd.Series(outlier_percentages).sort_values(ascending=False)

    table = doc.add_table(rows=1, cols=3)  # Added a column for outlier percentage
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Outliers'
    hdr_cells[2].text = 'Outlier %'  # New column for outlier percentage

    for col, percentage in outlier_percentage_series.head(args.top_n).items():
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = str(outlier_counts[col])
        row_cells[2].text = f'{percentage:.2f}%'  # Add percentage

    doc.add_heading('6. PCA Analysis', level=1)
    X = df[feature_cols].fillna(df[feature_cols].mean())
    X = StandardScaler().fit_transform(X)
    pca = PCA(n_components=args.pca_var)
    pca.fit(X)
    n_components = pca.n_components_
    doc.add_paragraph(f'Components for {args.pca_var * 100}% variance: {n_components}')

    # Transform data to get PCA components
    X_pca = pca.transform(X)

    # Create a DataFrame with PCA components
    pca_df = pd.DataFrame(
        X_pca[:, :min(args.top_n, n_components)],
        columns=[f'PC{i + 1}' for i in range(min(args.top_n, n_components))]
    )

    # Calculate correlation between PCA components and target
    pca_df[target_col] = df[target_col].values
    pca_target_corrs = pca_df.corrwith(pca_df[target_col]).drop(target_col).sort_values(ascending=False)

    # Add PCA-Target correlation table
    doc.add_paragraph(f'PCA Component Correlation with Target:')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Correlation with Target'

    for comp, corr in pca_target_corrs.items():
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = f'{corr:.3f}'

    # Plot PCA component correlations with target
    plt.figure(figsize=(10, 6))
    sns.barplot(x=pca_target_corrs.index, y=pca_target_corrs.values)
    plt.title(f'PCA Component Correlations with {target_col}')
    plt.xlabel('PCA Component')
    plt.ylabel(f'Correlation with {target_col}')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig('pca_target_correlation.png', dpi=300, bbox_inches='tight')
    plt.close()

    doc.add_paragraph('\nPCA Component Correlations with Target:')
    doc.add_picture('pca_target_correlation.png', width=Inches(6))

    # Show PCA component loadings
    loadings = pca.components_
    for i in range(min(3, n_components)):
        doc.add_paragraph(f'Component {i + 1} Top Loadings:')
        component_loadings = loadings[i]
        sorted_idx = np.argsort(np.abs(component_loadings))[::-1]
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Feature'
        hdr_cells[1].text = 'Loading'
        for idx in sorted_idx[:args.top_n]:
            row_cells = table.add_row().cells
            row_cells[0].text = feature_cols[idx]
            row_cells[1].text = f'{component_loadings[idx]:.3f}'

    # Plot scatterplots of top 2 PCA components vs target
    if n_components >= 2:
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))

        # PC1 vs Target
        sns.scatterplot(x=pca_df['PC1'], y=pca_df[target_col], alpha=0.6, ax=ax1)
        ax1.set_title(f'PC1 vs {target_col} (r = {pca_target_corrs["PC1"]:.2f})')
        ax1.set_xlabel('PC1')
        ax1.set_ylabel(target_col)

        # PC2 vs Target
        sns.scatterplot(x=pca_df['PC2'], y=pca_df[target_col], alpha=0.6, ax=ax2)
        ax2.set_title(f'PC2 vs {target_col} (r = {pca_target_corrs["PC2"]:.2f})')
        ax2.set_xlabel('PC2')
        ax2.set_ylabel(target_col)

        plt.tight_layout()
        plt.savefig('pca_top_components_vs_target.png', dpi=300, bbox_inches='tight')
        plt.close()

        doc.add_paragraph('\nTop PCA Components vs Target:')
        doc.add_picture('pca_top_components_vs_target.png', width=Inches(6))

    # PCA variance plot
    plt.figure(figsize=(8, 5))
    plt.plot(np.cumsum(pca.explained_variance_ratio_))
    plt.xlabel('Components')
    plt.ylabel('Explained Variance')
    plt.title('PCA Explained Variance')
    plt.grid(True, alpha=0.3)
    plt.savefig('pca_plot.png')
    doc.add_picture('pca_plot.png', width=Inches(6))
    plt.close()

    plt.figure(figsize=(8, 5))
    sns.histplot(df[target_col], kde=True)
    plt.title('Target Distribution')
    plt.savefig('target_dist.png')
    doc.add_picture('target_dist.png', width=Inches(6))
    plt.close()

    doc.add_heading('7. Target Variable Analysis', level=1)

    # Call the analysis function
    target_analysis = analyze_target_column(df, target_col, args.alpha)

    # Add target statistics to the report
    doc.add_paragraph('Target Variable Statistics:')
    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = 'Table Grid'
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = 'Statistic'
    hdr_cells[1].text = 'Value'

    for stat, value in target_analysis['statistics'].items():
        row_cells = stats_table.add_row().cells
        row_cells[0].text = stat.capitalize()
        row_cells[1].text = f'{value:.4f}' if isinstance(value, float) else str(value)

    # Add outlier information
    doc.add_paragraph('\nTarget Variable Outliers:')
    doc.add_paragraph(f"Outlier count: {target_analysis['outliers']['count']}")
    doc.add_paragraph(f"Outlier percentage: {target_analysis['outliers']['percentage']:.2f}%")
    doc.add_paragraph(f"Lower bound: {target_analysis['outliers']['lower_bound']:.4f}")
    doc.add_paragraph(f"Upper bound: {target_analysis['outliers']['upper_bound']:.4f}")

    # Add normality information
    doc.add_paragraph('\nTarget Variable Normality:')
    doc.add_paragraph(f"Shapiro-Wilk p-value: {target_analysis['normality']['p_value']:.4f}")
    doc.add_paragraph(f"Distribution: {target_analysis['normality']['interpretation']}")
    doc.add_paragraph(f"Skewness: {target_analysis['normality']['skewness']:.4f}")
    doc.add_paragraph(f"Kurtosis: {target_analysis['normality']['kurtosis']:.4f}")

    # Create a box plot to visualize outliers
    plt.figure(figsize=(8, 6))
    sns.boxplot(y=df[target_col])
    plt.title(f'Box Plot of {target_col}')
    plt.tight_layout()
    plt.savefig('target_boxplot.png')
    plt.close()
    doc.add_picture('target_boxplot.png', width=Inches(6))

    # Create Q-Q plot to assess normality
    plt.figure(figsize=(8, 6))
    stats.probplot(df[target_col].dropna(), plot=plt)
    plt.title(f'Q-Q Plot of {target_col}')
    plt.tight_layout()
    plt.savefig('target_qqplot.png')
    plt.close()
    doc.add_picture('target_qqplot.png', width=Inches(6))

    doc.save(args.report)

    # Clean up temporary image files
    for file in ['correlation_heatmap.png', 'pca_plot.png', 'target_dist.png',
                 'top_features_vs_target.png', 'pca_target_correlation.png',
                 'pca_top_components_vs_target.png', 'target_boxplot.png', 'target_qqplot.png']:
        if os.path.exists(file):
            os.remove(file)


if __name__ == '__main__':
    main()