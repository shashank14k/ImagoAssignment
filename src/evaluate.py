import numpy as np
import pandas as pd
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
import shap
from docx import Document
from docx.shared import Inches
import tempfile
from typing import Dict
import matplotlib.pyplot as plt
from src.infer import Predictor
from src.logger import logger


class Evaluator:
    def __init__(self, predictor: Predictor, r_square_thresh: float = 0.3):
        self.predictor = predictor
        self.doc = Document()
        self.doc.add_heading('Model Evaluation Report', 0)
        self.thresh = r_square_thresh

    def _add_metric_section(self, metrics: Dict, section_name: str):
        self.doc.add_heading(section_name, level=1)
        table = self.doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'RMSE'
        hdr_cells[1].text = 'MAE'
        hdr_cells[2].text = 'R²'
        row_cells = table.add_row().cells
        row_cells[0].text = f"{metrics['RMSE']:.4f}"
        row_cells[1].text = f"{metrics['MAE']:.4f}"
        row_cells[2].text = f"{metrics['R2']:.4f}"

    def _add_shap_plot(self, shap_values, features, feature_names, section_name: str):
        self.doc.add_heading(section_name, level=2)
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
            shap.summary_plot(shap_values[:, :, 0], features, feature_names=feature_names, show=False)
            plt.gcf().set_size_inches(8, 4)
            plt.savefig(tmpfile.name, bbox_inches='tight', dpi=300)
            plt.clf()
            self.doc.add_picture(tmpfile.name, width=Inches(6))

    def evaluate_metrics(self, df: pd.DataFrame) -> Dict[str, float]:
        y_true = df[self.predictor.config.target_col_name].values
        y_pred = np.array([self.predictor.infer(df.iloc[i].to_dict()) for i in range(len(df))])
        return {
            "RMSE": np.sqrt(mean_squared_error(y_true, y_pred)),
            "MAE": mean_absolute_error(y_true, y_pred),
            "R2": r2_score(y_true, y_pred)
        }

    def compute_shap_values(self, df: pd.DataFrame, sample_size: int = 100):
        df_sample = df.sample(min(sample_size, len(df)))
        X_sample = self.predictor.df_to_feature_tensor(df_sample).cpu()
        explainer = shap.DeepExplainer(self.predictor.model.cpu(), X_sample)
        return explainer.shap_values(X_sample), X_sample, explainer.expected_value

    def create_report(self, train_df: pd.DataFrame, test_df: pd.DataFrame,
                      output_path: str = "evaluation_report.docx", shap_sample_size: int = 100):
        train_metrics = self.evaluate_metrics(train_df)
        test_metrics = self.evaluate_metrics(test_df)
        if test_metrics["R2"] < self.thresh:
            logger.warning(
                "------------Warning R2 value={} less than {} on test set. Model might not be fit for deployment------------".format(
                    test_metrics["R2"], self.thresh))
        self._add_metric_section(train_metrics, "Training Metrics")
        self._add_metric_section(test_metrics, "Test Metrics")

        try:
            train_shap, train_data, train_base = self.compute_shap_values(train_df, shap_sample_size)
            test_shap, test_data, test_base = self.compute_shap_values(test_df, shap_sample_size)

            feature_names = self.predictor.feature_list
            self._add_shap_plot(train_shap, train_data, feature_names, "Training SHAP Summary")
            self.doc.add_paragraph("BASE SHAP VALUE {}".format(train_base))
            self._add_shap_plot(test_shap, test_data, feature_names, "Test SHAP Summary")
            self.doc.add_paragraph("BASE SHAP VALUE {}".format(test_base))
        except Exception as e:
            logger.warning("-----------Failed to generate SHAP values for model {}-----------------".format(
                self.predictor.config.model_params.name))
            logger.error("Error which occured {}".format(e))
        self._add_mir_section(train_df, "Train MIR", len(self.predictor.feature_list))
        self.doc.save(output_path)
        return output_path

    def _add_mir_section(self, df: pd.DataFrame, section_name: str, top_n: int = 10):
        """
        Computes Mutual Information Regression (MIR) scores for the features relative to the target,
        plots the feature importances, and adds the plot to the document.
        """
        self.doc.add_heading(section_name, level=2)
        from sklearn.feature_selection import mutual_info_regression
        X = df[self.predictor.feature_list]
        y = df[self.predictor.config.target_col_name].values
        mi_scores = mutual_info_regression(X, y, random_state=42)
        mi_scores_series = pd.Series(mi_scores, index=self.predictor.feature_list)
        mi_scores_series = mi_scores_series.sort_values(ascending=False)
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Feature'
        hdr_cells[1].text = 'MI Score'
        for feature, score in mi_scores_series.head(top_n).items():
            row_cells = table.add_row().cells
            row_cells[0].text = feature
            row_cells[1].text = f"{score:.4f}"
        import tempfile
        plt.figure(figsize=(8, 4))
        mi_scores_series.plot(kind='bar', color='skyblue')
        plt.ylabel("Mutual Information Score")
        plt.title("Feature Importance (Mutual Information)")
        plt.tight_layout()
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
            plt.savefig(tmpfile.name, bbox_inches='tight', dpi=300)
            plt.clf()
            self.doc.add_picture(tmpfile.name, width=Inches(6))


def main(args):
    logger.info("Starting evaluation report generation...")
    predictor = Predictor(args.config, args.checkpoint, device=args.device)
    evaluator = Evaluator(predictor)
    train_df = pd.read_csv(args.train)
    test_df = pd.read_csv(args.test)
    report_path = evaluator.create_report(train_df, test_df, args.report)
    logger.info(f"Report generated: {report_path}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Evaluate model and generate a report")
    parser.add_argument("--config", type=str, required=True,
                        help="Path to the config file (e.g., configs/conf.yaml)")
    parser.add_argument("--checkpoint", type=str, required=True,
                        help="Path to the checkpoint file (e.g., logs/1/best_model.pth)")
    parser.add_argument("--train", type=str, required=True,
                        help="Path to the training CSV file")
    parser.add_argument("--test", type=str, required=True,
                        help="Path to the test CSV file")
    parser.add_argument("--report", type=str, default="model_evaluation2.docx",
                        help="Output report file path (default: model_evaluation2.docx)")
    parser.add_argument("--device", type=str, help="device name cuda/cpu/mps")
    args = parser.parse_args()
    main(args)
